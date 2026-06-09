"""Generate formatted DOCX and PDF resumes with multiple template styles."""

import io
import os
from typing import Optional, List, Tuple

# DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import re

# PDF
from fpdf import FPDF

from app.services.templates import TemplateStyle, get_template, TemplateColors


def _make_safe_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in " _-." else "_" for c in name)


def _section_keywords() -> List[str]:
    """Known section header keywords for detection."""
    return [
        "professional summary", "summary", "profile", "objective", "about me",
        "experience", "work experience", "employment", "work history", "professional experience",
        "skills", "technical skills", "core competencies", "expertise",
        "education", "academic background",
        "certifications", "certificates", "licenses",
        "projects", "personal projects",
        "publications", "research",
        "languages", "interests", "activities",
        "references", "volunteer", "volunteering",
        "awards", "honors", "achievements",
        "leadership", "affiliations", "memberships",
    ]


def _normalize_section_name(text: str) -> str:
    """Normalize a section name for matching."""
    return text.lower().strip().rstrip(":").strip()


def _is_section_header(para) -> bool:
    """Check if a paragraph looks like a resume section header."""
    text = para.text.strip()
    if not text or len(text) > 60:
        return False

    lower = text.lower()

    # Direct keyword match (most reliable)
    if any(kw in lower for kw in _section_keywords()):
        return True

    # ALL-CAPS heuristic — only if it's short and looks like a section heading
    words = text.split()
    if text == text.upper() and len(text) > 2 and len(text) < 45 and len(words) <= 4:
        # Only match if it contains at least one section-like word
        section_words = {"summary", "experience", "education", "skills", "projects",
                        "certifications", "publications", "references", "experience",
                        "objective", "profile", "interests", "languages", "awards",
                        "volunteer", "leadership", "affiliations", "achievements"}
        if any(w.lower() in section_words for w in words) or len(words) <= 2:
            return True

    return False


def _identify_docx_sections(paragraphs) -> List[Tuple[int, int]]:
    """Identify section header positions and their content boundaries.
    Returns list of (header_p_idx, end_p_idx) tuples.
    """
    headers = []
    for i, para in enumerate(paragraphs):
        if _is_section_header(para):
            headers.append(i)

    if not headers:
        return []

    sections = []
    for idx, hpos in enumerate(headers):
        end = headers[idx + 1] if idx + 1 < len(headers) else len(paragraphs)
        sections.append((hpos, end))

    return sections


def _parse_tailored_into_sections(tailored_text: str) -> dict:
    """Parse tailored resume text into {section_name: [content_lines]}.

    Detects section headers heuristically:
    - Lines that are ALL CAPS with length > 2
    - Lines starting/ending with === markers
    - Lines matching known section keywords
    """
    sections = {}
    current_section = None
    current_lines = []

    # Lines to skip (visual separators, page breaks, etc.)
    skip_patterns = [
        lambda s: s.startswith("_") and len(s) > 10,  # underline separators
        lambda s: s.startswith("=") and len(s) > 10,  # equals separators
        lambda s: s.startswith("-") and len(s) > 10,  # dash separators
        lambda s: "PAGE BREAK" in s.upper(),  # page break markers
    ]

    def should_skip(line: str) -> bool:
        return any(pattern(line) for pattern in skip_patterns)

    for line in tailored_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if should_skip(stripped):
            continue

        # Detect section header
        is_header = False
        clean = stripped

        if stripped.startswith("=") or stripped.startswith("---"):
            clean = stripped.strip("= ").strip("-")
            is_header = True
        elif stripped == stripped.upper() and len(stripped) > 2 and len(stripped) < 50:
            # Check it's not just a normal uppercase line (like an acronym)
            # Section headers are typically short phrases
            words = stripped.split()
            if any(len(w) > 1 for w in words):  # avoid single-char
                maybe = stripped.lower()
                # Common resume section names
                known = _section_keywords()
                if any(kw in maybe for kw in known) or (stripped.isupper() and len(words) <= 4):
                    is_header = True

        if is_header:
            if current_section and current_lines:
                sections[current_section] = current_lines
            current_section = clean
            current_lines = []
        elif current_section:
            current_lines.append(stripped)

    # Don't forget the last section
    if current_section and current_lines:
        sections[current_section] = current_lines

    return sections


def _find_best_section_match(header_text: str, doc_sections: List[Tuple[int, int]], paragraphs) -> Optional[int]:
    """Find the best matching section in the document for a given header."""
    normalized = _normalize_section_name(header_text)
    best_score = 0
    best_idx = None

    for idx, (hpos, _) in enumerate(doc_sections):
        doc_header = paragraphs[hpos].text.strip()
        doc_norm = _normalize_section_name(doc_header)

        # Exact match
        if doc_norm == normalized:
            return idx

        # Check if one contains the other
        score = 0
        if normalized in doc_norm or doc_norm in normalized:
            score = len(min(normalized, doc_norm, key=len)) / len(max(normalized, doc_norm, key=len))

        # Check keyword overlap
        keywords = set(normalized.split())
        doc_keywords = set(doc_norm.split())
        overlap = keywords & doc_keywords
        if overlap:
            score = max(score, len(overlap) / max(len(keywords), len(doc_keywords), 1))

        if score > best_score:
            best_score = score
            best_idx = idx

    return best_idx if best_score > 0.3 else None


def generate_from_original_docx(
    original_file_path: str,
    tailored_text: str,
    candidate_name: str = "",
    job_title: str = "",
) -> bytes:
    """Generate a tailored resume using the original DOCX as a template.

    Opens the original uploaded DOCX, finds section headers, and replaces
    content text within each section while preserving ALL original formatting
    (fonts, colors, sizes, spacing, layout).

    Falls back to the regular template-based generation if the original
    file is not a DOCX or cannot be processed.
    """
    if not original_file_path or not os.path.exists(original_file_path):
        raise FileNotFoundError(f"Original resume file not found: {original_file_path}")

    ext = os.path.splitext(original_file_path)[1].lower()
    if ext != ".docx":
        raise ValueError(f"Original template generation only supports DOCX files, got: {ext}")

    doc = Document(original_file_path)
    paragraphs = doc.paragraphs

    # 1. Identify section boundaries in the original document
    doc_sections = _identify_docx_sections(paragraphs)

    if not doc_sections:
        # Fall back: no sections found, just return a template-based version
        from app.services.templates import get_template
        raise ValueError("Could not identify sections in the original document")

    # 2. Parse the tailored text into sections
    tailored_sections = _parse_tailored_into_sections(tailored_text)

    if not tailored_sections:
        raise ValueError("Could not parse tailored text into sections")

    # 3. For each tailored section, find the matching section in the original doc
    # and replace content while preserving formatting
    for header_text, content_lines in tailored_sections.items():
        match_idx = _find_best_section_match(header_text, doc_sections, paragraphs)
        if match_idx is None:
            continue

        hpos, end_pos = doc_sections[match_idx]
        section_start = hpos + 1  # first content paragraph after header
        section_end = end_pos

        # The number of paragraphs we have to work with in the original
        available_paras = section_end - section_start
        if available_paras <= 0:
            continue

        # Combine content lines into a single text block
        new_content = "\n".join(content_lines)

        if available_paras == 1:
            # Single paragraph — replace all runs with new content
            para = paragraphs[section_start]
            _replace_paragraph_text(para, new_content)
        else:
            # Multiple paragraphs — distribute content across them
            _distribute_content_across_paragraphs(
                paragraphs, section_start, section_end, content_lines
            )

    # 4. Return the modified document
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _replace_paragraph_text(para, new_text: str):
    """Replace all text in a paragraph's runs, preserving formatting of the first run."""
    if not para.runs:
        # No runs — create one
        run = para.add_run(new_text)
    else:
        # Set text in first run, clear others
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


def _distribute_content_across_paragraphs(
    paragraphs, start: int, end: int, content_lines: List[str]
):
    """Distribute content lines across available paragraphs preserving formatting."""
    available = end - start

    if available <= 0:
        return

    if len(content_lines) == 0:
        # No content — clear all paragraphs
        for i in range(start, end):
            _replace_paragraph_text(paragraphs[i], "")
        return

    # Simple distribution: group lines evenly across available paragraphs
    lines_per_para = max(1, (len(content_lines) + available - 1) // available)

    for i in range(available):
        para_idx = start + i
        if para_idx >= len(paragraphs):
            break

        line_start = i * lines_per_para
        line_end = min((i + 1) * lines_per_para, len(content_lines))

        if line_start < len(content_lines):
            group = content_lines[line_start:line_end]
            _replace_paragraph_text(paragraphs[para_idx], "\n".join(group))
        else:
            _replace_paragraph_text(paragraphs[para_idx], "")


def _rgb(triple):
    return RGBColor(*triple)


def _hex_from_rgb(triple):
    """Convert RGB tuple to fpdf-compatible string."""
    return "#{:02x}{:02x}{:02x}".format(*triple)


# ── DOCX TEMPLATE RENDERERS ──────────────────────────────────────────────

def _docx_section_underline(doc, section: str, style: TemplateStyle):
    c = _rgb(style.colors.primary)
    p = doc.add_paragraph()
    r = p.add_run(section)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = c
    r.font.name = style.font_heading
    # Underline border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "{:02x}{:02x}{:02x}".format(*style.colors.accent),
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_section_bar(doc, section: str, style: TemplateStyle):
    # Left bar + heading
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    left = pBdr.makeelement(qn("w:left"), {
        qn("w:val"): "single",
        qn("w:sz"): "24",
        qn("w:space"): "8",
        qn("w:color"): "{:02x}{:02x}{:02x}".format(*style.colors.accent),
    })
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(section)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = _rgb(style.colors.primary)
    r.font.name = style.font_heading


def _docx_section_minimal(doc, section: str, style: TemplateStyle):
    p = doc.add_paragraph()
    r = p.add_run(section)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = _rgb(style.colors.primary)
    r.font.name = style.font_heading


def _docx_section_badge(doc, section: str, style: TemplateStyle):
    # Badge-style: colored background highlight
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "{:02x}{:02x}{:02x}".format(*style.colors.accent),
    })
    pPr.append(shd)
    # Add spacing
    pPr.append(pPr.makeelement(qn("w:spacing"), {
        qn("w:before"): "120",
        qn("w:after"): "120",
    }))
    r = p.add_run(f"  {section}  ")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.name = style.font_heading


SECTION_RENDERERS_DOCX = {
    "underline": _docx_section_underline,
    "bar": _docx_section_bar,
    "minimal": _docx_section_minimal,
    "badge": _docx_section_badge,
}


def generate_docx(
    tailored_text: str,
    candidate_name: str = "Candidate",
    job_title: str = "",
    template_name: str = "professional",
    contact_info: str = "",
) -> bytes:
    """Generate a formatted .docx resume using the selected template."""
    style = get_template(template_name)
    doc = Document()

    # ── Page margins (fit more content per page) ──
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── Default style ──
    normal = doc.styles["Normal"]
    normal.font.name = style.font_body
    normal.font.size = Pt(10)
    normal.font.color.rgb = _rgb(style.colors.text)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    # ── Name (18pt matching original PDF) ──
    align = WD_ALIGN_PARAGRAPH.CENTER if style.header_align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    h = doc.add_heading(candidate_name, level=0)
    h.alignment = align
    for run in h.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = _rgb(style.colors.primary)
        run.font.name = style.font_heading

    # ── Contact info line (matching original PDF: location | email | phone | linkedin) ──
    if contact_info:
        p = doc.add_paragraph()
        p.alignment = align
        r = p.add_run(contact_info)
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb(style.colors.text)
        r.font.name = style.font_body
        p.paragraph_format.space_after = Pt(6)
    elif job_title:
        # Fallback: show job title as subtitle if no contact info
        p = doc.add_paragraph()
        p.alignment = align
        r = p.add_run(f"Tailored for: {job_title}")
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(style.colors.accent)
        r.italic = False
        r.font.name = style.font_body
        p.paragraph_format.space_after = Pt(6)
    else:
        # Add small space
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

    # ── Parse sections using shared logic ──
    sections = _parse_tailored_into_sections(tailored_text)

    # Define section order (matches PDF)
    section_order = [
        "summary", "professional summary", "profile", "objective",
        "skills", "technical skills", "core competencies", "expertise",
        "experience", "work experience", "employment", "work history",
        "education", "academic background",
        "certifications", "certificates", "licenses",
        "additional information",
        "other projects", "personal projects", "projects",
        "publications", "research",
        "languages", "interests", "activities",
        "references", "volunteer", "volunteering",
        "awards", "honors", "achievements",
        "leadership", "affiliations", "memberships",
    ]

    section_renderer = SECTION_RENDERERS_DOCX.get(style.section_style, _docx_section_underline)
    rendered_sections = set()

    for section_name in section_order:
        for key in sections:
            if section_name in key.lower() and key not in rendered_sections:
                section_renderer(doc, key, style)
                for line in sections[key]:
                    line = line.strip()
                    if not line:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                        continue
                    if line.startswith("●") or line.startswith("-") or line.startswith("*"):
                        p = doc.add_paragraph(line[1:].strip())
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.left_indent = Inches(0.25)
                    else:
                        p = doc.add_paragraph(line)
                        p.paragraph_format.space_after = Pt(1)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                rendered_sections.add(key)
                break

    # Render any remaining sections
    for key, content in sections.items():
        if key not in rendered_sections:
            section_renderer(doc, key, style)
            for line in content:
                line = line.strip()
                if not line:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    continue
                if line.startswith("●") or line.startswith("-") or line.startswith("*"):
                    p = doc.add_paragraph(line[1:].strip())
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.left_indent = Inches(0.25)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PDF Text Sanitization ──────────────────────────────────────────────


def _sanitize_pdf_text(text: str) -> str:
    """Replace characters not supported by Helvetica (Latin-1) with ASCII equivalents."""
    replacements = {
        "\u2022": "\u00B7",  # bullet → middle dot (Latin-1)
        "\u2023": ">",       # triangular bullet → >
        "\u25E6": "\u00B7",  # white bullet → middle dot
        "\u25CF": "\u00B7",  # black circle → middle dot
        "\u25CB": "o",       # white circle → o
        "\u25A0": "\u00B7",  # black square → middle dot
        "\u25AA": "\u00B7",  # black small square → middle dot
        "\u2605": "*",       # star → *
        "\u2606": "*",       # white star → *
        "\u2713": "[x]",     # check mark → [x]
        "\u2717": "[ ]",     # ballot x → [ ]
        "\u2013": "-",       # en dash → -
        "\u2014": "--",      # em dash → --
        "\u2018": "'",       # left single quote → '
        "\u2019": "'",       # right single quote → '
        "\u201C": '"',       # left double quote → "
        "\u201D": '"',       # right double quote → "
        "\u2026": "...",     # ellipsis → ...
        "\u25B6": ">",       # play button → >
        "\u25C0": "<",       # reverse button → <
        "\u203A": ">",       # single right angle quote → >
        "\u2039": "<",       # single left angle quote → <
        "\u00AB": "<<",      # double left angle -> <<
        "\u00BB": ">>",      # double right angle -> >>
        "\u00A0": " ",        # nbsp → space
        "\u200B": "",         # zero-width space → empty
        "\u2009": " ",        # thin space → space
        "\u00A9": "(c)",     # copyright
        "\u00AE": "(r)",     # registered
        "\u2122": "(tm)",    # trademark
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text


def _pdf_section_underline(pdf, section: str, style: TemplateStyle):
    is_naukri = style.name == "naukri"
    # Section headers - naukri uses ALL CAPS, slightly larger, accent color underline
    font_size = 11 if is_naukri else 10
    pdf.set_font("Helvetica", "B", font_size)
    pdf.set_text_color(*style.colors.primary)
    # Apply header alignment to section headers for consistency
    align = "C" if style.header_align == "center" else "L"
    section_text = section.upper() if is_naukri else section
    pdf.cell(0, 6, section_text, new_x="LMARGIN", new_y="NEXT", align=align)
    pdf.ln(0.5 if is_naukri else 1)
    # Full-width underline in ACCENT color (naukri uses blue accent line)
    underline_color = style.colors.accent if is_naukri else style.colors.primary
    pdf.set_draw_color(*underline_color)
    pdf.set_line_width(0.5 if is_naukri else 0.4)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(2.5 if is_naukri else 3.5)
    pdf.set_text_color(*style.colors.text)
    pdf.set_line_width(0.2)  # reset to default


def _pdf_section_bar(pdf, section: str, style: TemplateStyle):
    # Draw left bar
    x, y = pdf.get_x(), pdf.get_y()
    pdf.set_fill_color(*style.colors.accent)
    pdf.rect(x, y, 2, 7, style="F")
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*style.colors.primary)
    pdf.cell(6)  # indent after bar
    # Apply header alignment to section headers for consistency
    align = "C" if style.header_align == "center" else "L"
    pdf.cell(0, 6, section, new_x="LMARGIN", new_y="NEXT", align=align)
    pdf.ln(3)
    pdf.set_text_color(*style.colors.text)


def _pdf_section_minimal(pdf, section: str, style: TemplateStyle):
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*style.colors.primary)
    # Apply header alignment to section headers for consistency
    align = "C" if style.header_align == "center" else "L"
    pdf.cell(0, 6, section, new_x="LMARGIN", new_y="NEXT", align=align)
    pdf.ln(3)
    pdf.set_text_color(*style.colors.text)


def _pdf_section_badge(pdf, section: str, style: TemplateStyle):
    pdf.set_font("Helvetica", "B", 10)
    # Colored background
    pdf.set_fill_color(*style.colors.accent)
    pdf.set_text_color(255, 255, 255)
    text_w = pdf.get_string_width(section) + 8

    if style.header_align == "center":
        # Calculate position to center the badge
        total_w = text_w + 6  # 3 units padding on each side
        x = (pdf.w - total_w) / 2
        pdf.set_x(x)
        pdf.cell(text_w, 7, f"  {section}  ", fill=True, new_x="LMARGIN", new_y="NEXT")
    else:
        # Left alignment - start from left margin with padding
        pdf.cell(3)  # left padding
        pdf.cell(text_w, 7, f"  {section}  ", fill=True, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(4)
    pdf.set_text_color(*style.colors.text)


SECTION_RENDERERS_PDF = {
    "underline": _pdf_section_underline,
    "bar": _pdf_section_bar,
    "minimal": _pdf_section_minimal,
    "badge": _pdf_section_badge,
}


class ResumePDF(FPDF):
    def __init__(self, style: TemplateStyle):
        super().__init__()
        self.style = style
        # Naukri template uses tighter margins for more content space
        if style.name == "naukri":
            self.set_margins(15, 15, 15)
            self.set_auto_page_break(auto=True, margin=20)
        else:
            # Use slightly larger margins for better visual appeal
            self.set_margins(18, 18, 18)  # Increased from 15 to 18
            self.set_auto_page_break(auto=True, margin=30)  # Increased from 25 to 30

    def header(self):
        pass

    def footer(self):
        # No footer — matches traditional resume format (no branding watermark)
        pass


def generate_pdf(
    tailored_text: str,
    candidate_name: str = "Candidate",
    job_title: str = "",
    template_name: str = "naukri",
    contact_info: str = "",
) -> bytes:
    """Generate a formatted PDF resume using the selected template.

    ``contact_info`` is an optional string like "Delhi, India | email@example.com | 123-456-7890"
    rendered below the candidate name, matching traditional resume layout.
    """
    style = get_template(template_name)
    pdf = ResumePDF(style)
    pdf.add_page()

    is_naukri = style.name == "naukri"

    # ── Name ──
    name_size = 20 if is_naukri else 18
    pdf.set_font("Helvetica", "B", name_size)
    pdf.set_text_color(*style.colors.primary)
    name_safe = _sanitize_pdf_text(candidate_name)
    if style.header_align == "center":
        pdf.cell(0, 11 if is_naukri else 11, name_safe, new_x="LMARGIN", new_y="NEXT", align="C")
    else:
        pdf.cell(0, 11, name_safe, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2 if is_naukri else 3)

    # ── Contact info line (matching original PDF: location | email | phone | linkedin) ──
    if contact_info:
        pdf.set_font("Helvetica", "", 9 if is_naukri else 10)
        pdf.set_text_color(*style.colors.secondary_text)
        info_safe = _sanitize_pdf_text(contact_info)
        if style.header_align == "center":
            pdf.cell(0, 5, info_safe, new_x="LMARGIN", new_y="NEXT", align="C")
        else:
            pdf.cell(0, 5, info_safe, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4 if is_naukri else 5)
    elif job_title:
        # Fallback: show job title as subtitle if no contact info
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*style.colors.accent)
        title_safe = _sanitize_pdf_text(job_title)
        if style.header_align == "center":
            pdf.cell(0, 5, f"Tailored for: {title_safe}", new_x="LMARGIN", new_y="NEXT", align="C")
        else:
            pdf.cell(0, 5, f"Tailored for: {title_safe}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
    else:
        pdf.ln(2)

    # ── Body ──
    pdf.set_text_color(*style.colors.text)

    # Parse the tailored text into sections first
    sections = _parse_tailored_into_sections(tailored_text)

    # If parsing failed, fall back to line-by-line processing with better logic
    if not sections:
        _generate_pdf_line_by_line(pdf, tailored_text, candidate_name, contact_info, style)
    else:
        _generate_pdf_from_sections(pdf, sections, style)

    return bytes(pdf.output())


def _generate_pdf_line_by_line(pdf, tailored_text: str, candidate_name: str, contact_info: str, style: TemplateStyle):
    """Fallback method: generate PDF line by line with improved logic."""
    section_renderer = SECTION_RENDERERS_PDF.get(style.section_style, _pdf_section_underline)

    _email_re = re.compile(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}')
    _phone_re = re.compile(r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')

    lines = tailored_text.split("\n")
    i = 0
    in_section = False

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines at the beginning
        if not line and i < 3:
            i += 1
            continue

        # Skip duplicate name/contact info
        if (line == candidate_name or
            (contact_info and line == contact_info) or
            (candidate_name and re.search(r'\b' + re.escape(candidate_name) + r'\b', line) and
             (_email_re.search(line) or _phone_re.search(line)))):
            i += 1
            continue

        # Check for section headers
        if (line.startswith("=") or line.startswith("---") or
            (line.isupper() and len(line) > 2 and len(line) < 50)):

            # Extract section header
            if line.startswith("=") or line.startswith("---"):
                section_name = line.strip("= -").strip()
            else:
                section_name = line

            # Render section header with proper spacing
            if in_section:
                pdf.ln(2)  # Add space before new section
            section_renderer(pdf, _sanitize_pdf_text(section_name), style)
            in_section = True
            i += 1
        elif line:
            # Regular content line
            is_naukri = style.name == "naukri"
            body_font_size = 9 if is_naukri else 10
            line_height = 5.0 if is_naukri else 5.5
            bullet_spacing = 0.3 if is_naukri else 0.5
            empty_line_spacing = 1 if is_naukri else 1.5

            pdf.set_font("Helvetica", "", body_font_size)
            pdf.set_text_color(*style.colors.text)

            # Handle bullet points with better formatting
            if line.startswith("●") or line.startswith("-") or line.startswith("*"):
                pdf.set_font("Helvetica", "", body_font_size)
                # Add proper bullet character (middle dot, Latin-1 compatible) and spacing
                bullet_text = "\u00B7 " + line[1:].strip()
                pdf.multi_cell(0, line_height, _sanitize_pdf_text(bullet_text))
                pdf.ln(bullet_spacing)
            else:
                # Regular text with proper spacing
                pdf.multi_cell(0, line_height, _sanitize_pdf_text(line))
                pdf.ln(bullet_spacing)
            in_section = True
            i += 1
        else:
            # Empty line with consistent spacing
            empty_line_spacing = 1 if style.name == "naukri" else 1.5
            pdf.ln(empty_line_spacing)
            i += 1


def _generate_pdf_from_sections(pdf, sections: dict, style: TemplateStyle):
    """Generate PDF from parsed sections with proper formatting."""
    section_renderer = SECTION_RENDERERS_PDF.get(style.section_style, _pdf_section_underline)

    # Define section order (common resume sections) - matches original resume structure
    # More specific section names first to avoid partial matches (e.g., "other projects" before "projects")
    section_order = [
        "summary", "professional summary", "profile", "objective",
        "skills", "technical skills", "core competencies", "expertise",
        "experience", "work experience", "employment", "work history",
        "education", "academic background",
        "certifications", "certificates", "licenses",
        "additional information",
        "other projects", "personal projects", "projects",
        "publications", "research",
        "languages", "interests", "activities",
        "references", "volunteer", "volunteering",
        "awards", "honors", "achievements",
        "leadership", "affiliations", "memberships",
    ]

    # Render sections in order
    rendered_sections = set()

    for section_name in section_order:
        for key in sections:
            if section_name in key.lower() and key not in rendered_sections:
                # Render section header
                section_renderer(pdf, _sanitize_pdf_text(key), style)

                # Render section content
                for line in sections[key]:
                    line = line.strip()
                    if not line:
                        pdf.ln(1.5)  # Consistent spacing for empty lines
                        continue

                    is_naukri = style.name == "naukri"
                body_font_size = 9 if is_naukri else 10
                line_height = 5.0 if is_naukri else 5.5
                bullet_spacing = 0.3 if is_naukri else 0.5
                section_spacing = 2 if is_naukri else 3
                empty_line_spacing = 1 if is_naukri else 1.5

                # Render section content
                for line in sections[key]:
                    line = line.strip()
                    if not line:
                        pdf.ln(empty_line_spacing)
                        continue

                    # Handle bullet points with better spacing
                    if line.startswith("●") or line.startswith("-") or line.startswith("*"):
                        pdf.set_font("Helvetica", "", body_font_size)
                        bullet_text = "\u00B7 " + line[1:].strip()
                        pdf.multi_cell(0, line_height, _sanitize_pdf_text(bullet_text))
                        pdf.ln(bullet_spacing)
                    else:
                        pdf.set_font("Helvetica", "", body_font_size)
                        # Handle long lines with automatic word wrap
                        pdf.multi_cell(0, line_height, _sanitize_pdf_text(line))
                        pdf.ln(bullet_spacing)

                # Add proper spacing after section
                pdf.ln(section_spacing)
                rendered_sections.add(key)
                break

    # Render any remaining sections that weren't in the predefined order
    for key, content in sections.items():
        if key not in rendered_sections:
            section_renderer(pdf, _sanitize_pdf_text(key), style)

            is_naukri = style.name == "naukri"
            body_font_size = 9 if is_naukri else 10
            line_height = 5.0 if is_naukri else 5.5
            bullet_spacing = 0.3 if is_naukri else 0.5
            section_spacing = 2 if is_naukri else 3
            empty_line_spacing = 1 if is_naukri else 1.5

            for line in content:
                line = line.strip()
                if not line:
                    pdf.ln(empty_line_spacing)
                    continue

                if line.startswith("●") or line.startswith("-") or line.startswith("*"):
                    pdf.set_font("Helvetica", "", body_font_size)
                    bullet_text = "\u00B7 " + line[1:].strip()
                    pdf.multi_cell(0, line_height, _sanitize_pdf_text(bullet_text))
                    pdf.ln(bullet_spacing)
                else:
                    pdf.set_font("Helvetica", "", body_font_size)
                    pdf.multi_cell(0, line_height, _sanitize_pdf_text(line))
                    pdf.ln(bullet_spacing)

            pdf.ln(section_spacing)
